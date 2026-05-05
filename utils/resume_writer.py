from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from pathlib import Path


# Section header keywords for detection
SECTION_KEYWORDS = {
    "summary", "objective", "profile", "about",
    "experience", "employment", "work history",
    "education", "academic",
    "skills", "technical skills", "core competencies",
    "projects", "certifications", "awards", "publications",
    "languages", "interests", "volunteer",
}


def _add_horizontal_rule(paragraph):
    """Add a bottom border line under a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E5F8A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _is_section_header(line: str) -> bool:
    """Detect if a line is a section header."""
    stripped = line.strip().rstrip(":").lower()
    if stripped in SECTION_KEYWORDS:
        return True
    if line.isupper() and 3 <= len(line.split()) <= 5:
        return True
    # Common patterns: "WORK EXPERIENCE", "TECHNICAL SKILLS", etc.
    if re.match(r"^[A-Z][A-Z\s/&]+$", line.strip()) and len(line.strip()) < 40:
        return True
    return False


def _is_contact_line(line: str) -> bool:
    return bool(re.search(r"@|\d{3}[-.\s]\d{4}|linkedin\.com|github\.com|http", line, re.I))


def create_tailored_docx(tailored_text: str, job: dict, output_path: str, candidate_name: str = "") -> str:
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    lines = [l.rstrip() for l in tailored_text.splitlines()]
    # Remove empty leading/trailing lines
    while lines and not lines[0].strip():
        lines.pop(0)

    first_content_line = True
    prev_was_blank = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            if not prev_was_blank:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
            prev_was_blank = True
            continue

        prev_was_blank = False

        # First non-blank line → Candidate name (large, centered)
        if first_content_line:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            run.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            first_content_line = False
            continue

        # Contact info line (email, phone, LinkedIn) → small, centered
        if _is_contact_line(stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            continue

        # Section header
        if _is_section_header(stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped.upper().rstrip(":"))
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x2E, 0x5F, 0x8A)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            _add_horizontal_rule(p)
            continue

        # Bullet point
        if stripped.startswith(("•", "-", "*", "◦", "▪")):
            content = stripped.lstrip("•-*◦▪ ").strip()
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(content)
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            continue

        # Bold line detection: wrapped in ** or all caps short line
        bold_match = re.match(r"^\*\*(.+?)\*\*$", stripped)
        if bold_match or (stripped.isupper() and len(stripped) < 60 and not _is_section_header(stripped)):
            content = bold_match.group(1) if bold_match else stripped
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(2)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)

    doc.save(output_path)
    return output_path


def create_cover_letter_docx(cover_text: str, job: dict, output_path: str, candidate_name: str = "") -> str:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Date
    from datetime import datetime
    p = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.name = "Calibri"
    p.paragraph_format.space_after = Pt(12)

    # Greeting / body
    for line in cover_text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue
        p = doc.add_paragraph(stripped)
        p.runs[0].font.size = Pt(10.5)
        p.runs[0].font.name = "Calibri"
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(14)

    doc.save(output_path)
    return output_path
