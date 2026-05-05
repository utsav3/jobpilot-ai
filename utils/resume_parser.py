import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
import re


def parse_resume(file_path: str) -> dict:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        text = _extract_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        text = _extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported format: {ext}")
    return {
        "raw_text": text.strip(),
        "file_path": file_path,
        "format": ext,
        "name": Path(file_path).stem,
    }


def _extract_pdf(path: str) -> str:
    doc = fitz.open(path)
    return "\n".join(page.get_text() for page in doc)


def _extract_docx(path: str) -> str:
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def extract_name_from_resume(text: str) -> str:
    """Best-effort extraction of candidate name from first lines."""
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    for line in lines[:5]:
        # Skip lines that look like contact info
        if "@" in line or re.search(r"\d{3}[-.\s]\d{3}", line):
            continue
        if len(line.split()) <= 4 and line[0].isupper():
            return line
    return "Candidate"
